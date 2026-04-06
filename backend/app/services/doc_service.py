import io
import os
import uuid
from pathlib import Path
from typing import Optional
from sqlalchemy.ext.asyncio import AsyncSession
from app.config import get_settings
from app.models.document import Document, DocumentChunk
from app.services.llm_service import get_embedding

settings = get_settings()

CHUNK_SIZE = 1000
CHUNK_OVERLAP = 200


def _split_text(text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> list[str]:
    """Simple recursive character text splitter."""
    if len(text) <= chunk_size:
        return [text.strip()] if text.strip() else []

    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunk = text[start:end]
        chunks.append(chunk.strip())
        start += chunk_size - overlap
    return [c for c in chunks if c]


def _extract_text_from_pdf(content: bytes) -> str:
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(content))
    text_parts = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            text_parts.append(text)
    return "\n".join(text_parts)


def _extract_text_from_docx(content: bytes) -> str:
    import docx
    doc = docx.Document(io.BytesIO(content))
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def _extract_text_from_xlsx(content: bytes) -> str:
    import openpyxl
    wb = openpyxl.load_workbook(io.BytesIO(content), read_only=True, data_only=True)
    parts = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        parts.append(f"=== Аркуш: {sheet_name} ===")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c).strip() if c is not None else "" for c in row]
            line = " | ".join(c for c in cells if c)
            if line:
                parts.append(line)
        parts.append("")
    wb.close()
    return "\n".join(parts)


def _extract_text_from_csv(content: bytes) -> str:
    import csv
    text = content.decode("utf-8", errors="replace")
    reader = csv.reader(io.StringIO(text))
    lines = []
    for row in reader:
        cells = [c.strip() for c in row if c.strip()]
        if cells:
            lines.append(" | ".join(cells))
    return "\n".join(lines)


def _extract_text(content: bytes, mime_type: str) -> str:
    if mime_type == "application/pdf":
        return _extract_text_from_pdf(content)
    elif mime_type in (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",):
        return _extract_text_from_docx(content)
    elif mime_type in (
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "application/vnd.ms-excel",):
        return _extract_text_from_xlsx(content)
    elif mime_type == "text/csv":
        return _extract_text_from_csv(content)
    else:
        return content.decode("utf-8", errors="replace")


async def ingest_document(
    file_content: bytes,
    original_name: str,
    mime_type: str,
    uploader_id: int,
    db: AsyncSession,) -> Document:
    os.makedirs(settings.UPLOAD_DIR, exist_ok=True)

    ext = Path(original_name).suffix
    unique_filename = f"{uuid.uuid4().hex}{ext}"
    file_path = os.path.join(settings.UPLOAD_DIR, unique_filename)

    with open(file_path, "wb") as f:
        f.write(file_content)

    document = Document(
        filename=unique_filename,
        original_name=original_name,
        file_path=file_path,
        uploaded_by=uploader_id,
        size_bytes=len(file_content),
        mime_type=mime_type,)
    db.add(document)
    await db.flush()
    await db.refresh(document)
    text = _extract_text(file_content, mime_type)
    chunks_text = _split_text(text)
    for idx, chunk_text in enumerate(chunks_text):
        try:
            embedding = await get_embedding(chunk_text)
        except Exception:
            embedding = None
        chunk = DocumentChunk(
            document_id=document.id,
            content=chunk_text,
            embedding=embedding,
            chunk_index=idx,)
        db.add(chunk)
    await db.flush()
    return document

async def delete_document_files(file_path: str) -> None:
    try:
        if os.path.exists(file_path):
            os.remove(file_path)
    except OSError:
        pass
